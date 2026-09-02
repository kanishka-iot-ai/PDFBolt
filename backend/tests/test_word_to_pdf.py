from pathlib import Path
import docx
from backend.app.processors.word_to_pdf import WordToPdfProcessor
from backend.app.core.validation import validate_pdf_output

FIXTURES_DIR = Path(__file__).parent / "fixtures"


def test_word_to_pdf_conversion(tmp_path):
    """INVARIANT: Generates valid PDF from Word .docx with headings and tables."""
    docx_path = tmp_path / "sample_doc.docx"
    doc = docx.Document()
    doc.add_heading("PDFBolt Executive Summary", level=1)
    doc.add_paragraph("This document tests Word to PDF conversion high fidelity output.")
    
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Throughput"
    table.rows[1].cells[1].text = "100 MB/s"
    table.rows[2].cells[0].text = "Fidelity"
    table.rows[2].cells[1].text = "100%"
    
    doc.save(str(docx_path))
    
    processor = WordToPdfProcessor(job_id="test_word2pdf", work_dir=tmp_path)
    result = processor.run([docx_path], {})
    
    assert result.status == "COMPLETED"
    assert result.output_path.exists()
    assert result.output_path.stat().st_size > 0
    page_count = validate_pdf_output(result.output_path)
    assert page_count >= 1


def test_word_to_pdf_process_bytes(tmp_path):
    """INVARIANT: process_bytes accepts in-memory Word buffer and returns PDF bytes."""
    docx_path = tmp_path / "temp.docx"
    doc = docx.Document()
    doc.add_heading("Invoice Note", level=1)
    doc.add_paragraph("Thank you for your business with PDFBolt.")
    doc.save(str(docx_path))
    
    with open(docx_path, "rb") as f:
        doc_bytes = f.read()
        
    processor = WordToPdfProcessor(job_id="test_word2pdf_bytes", work_dir=tmp_path)
    pdf_bytes, filename, metrics = processor.process_bytes(doc_bytes, "invoice.docx")
    
    assert len(pdf_bytes) > 0
    assert filename.endswith(".pdf")
    assert metrics["format"] == "pdf"
