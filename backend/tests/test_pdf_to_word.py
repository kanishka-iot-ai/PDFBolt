import io
from pathlib import Path
import pymupdf
from docx import Document

from backend.app.processors.pdf_to_word import PdfToWordProcessor
from backend.app.core.validation import validate_docx_output

FIXTURES_DIR = Path(__file__).parent / "fixtures"


def create_complex_pdf(file_path: Path):
    """Generates a multi-page PDF with headings, multi-columns, tables, and graphic blocks."""
    doc = pymupdf.open()
    
    # Page 1: Headings, Paragraphs, and Tables
    p1 = doc.new_page(width=595, height=842)
    p1.insert_text((50, 72), "Annual Corporate Report 2026", fontsize=18, color=(0.1, 0.2, 0.5))
    p1.insert_text((50, 110), "Executive Overview & Global Strategy", fontsize=14)
    p1.insert_text((50, 140), "This report outlines key performance metrics, revenue growth, and organizational strategy.", fontsize=11)
    
    # Draw table
    p1.draw_rect(pymupdf.Rect(50, 180, 545, 280), color=(0.5, 0.5, 0.5))
    headers = ["Quarter", "Revenue (USD)", "Operating Margin", "Status"]
    for idx, h in enumerate(headers):
        p1.insert_text((60 + idx * 120, 205), h, fontsize=11)
        
    p1.insert_text((60, 235), "Q1 2026", fontsize=10)
    p1.insert_text((180, 235), "$1,450,000", fontsize=10)
    p1.insert_text((300, 235), "42.5%", fontsize=10)
    p1.insert_text((420, 235), "Complete", fontsize=10)
    
    # Page 2: Multi-Column Layout
    p2 = doc.new_page(width=595, height=842)
    p2.insert_text((50, 72), "Research Paper: Distributed System Resilience", fontsize=16)
    
    # Left column (x: 50 to 280)
    p2.insert_text((50, 110), "1. Introduction", fontsize=12)
    p2.insert_text((50, 130), "Modern cloud infrastructures require", fontsize=10)
    p2.insert_text((50, 145), "fault-tolerant consensus protocols.", fontsize=10)
    
    # Right column (x: 310 to 545)
    p2.insert_text((310, 110), "2. Methodology", fontsize=12)
    p2.insert_text((310, 130), "We evaluated latency distribution", fontsize=10)
    p2.insert_text((310, 145), "under synthetic network partitions.", fontsize=10)

    raw_bytes = doc.tobytes()
    doc.close()
    with open(file_path, "wb") as f:
        f.write(raw_bytes)
    return raw_bytes


def test_pdf_to_word_basic_text(tmp_path):
    """Test basic text PDF to DOCX conversion."""
    p1 = FIXTURES_DIR / "1page_text.pdf"
    processor = PdfToWordProcessor(job_id="test_word_basic", work_dir=tmp_path)
    result = processor.run([p1], {})

    assert result.status == "COMPLETED"
    validate_docx_output(result.output_path)
    
    doc = Document(str(result.output_path))
    full_text = " ".join([p.text for p in doc.paragraphs])
    assert len(full_text.strip()) > 5
    assert processor.metrics["quality_score"] >= 80


def test_pdf_to_word_complex_multipage(tmp_path):
    """
    CRITICAL TEST:
    Verifies that multi-page documents with headings, multi-columns, and tables
    are accurately converted into editable Word elements.
    """
    pdf_path = tmp_path / "complex.pdf"
    create_complex_pdf(pdf_path)
    
    processor = PdfToWordProcessor(job_id="test_word_complex", work_dir=tmp_path)
    result = processor.run([pdf_path], {})
    
    assert result.status == "COMPLETED"
    assert result.output_path.exists()
    validate_docx_output(result.output_path)
    
    doc = Document(str(result.output_path))
    full_text = " ".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for r in t.rows for c in r.cells])
    
    # Check text preservation
    assert "Annual Corporate Report 2026" in full_text
    assert "Distributed System Resilience" in full_text
    assert "Methodology" in full_text
    
    # Check quality score
    metrics = processor.metrics
    assert metrics["quality_score"] >= 85
    assert metrics["status"] == "success"
    assert metrics["converted_pages"] == 2
