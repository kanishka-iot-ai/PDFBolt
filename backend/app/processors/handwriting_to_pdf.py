import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from backend.app.processors.handwriting import HandwritingProcessor
from backend.app.models.handwriting import ExportFormat


class HandwritingDocumentGenerator:
    """Generates clean output documents (PDF, DOCX, TXT) from transcribed notes."""

    @classmethod
    def generate(
        cls,
        pages_text: List[str],
        title: str = "Digitized Handwriting",
        design: Optional[Any] = None,
        export_format: Union[ExportFormat, str] = ExportFormat.PDF
    ) -> bytes:
        fmt_str = export_format.value if hasattr(export_format, "value") else str(export_format).lower()

        if fmt_str == "docx":
            return cls.generate_docx(pages_text, title)
        elif fmt_str == "txt":
            return cls.generate_txt(pages_text, title)
        else: # pdf
            return cls.generate_pdf(pages_text, title, design)

    @staticmethod
    def generate_pdf(
        pages_text: Union[List[str], str],
        title: str = "Digitized Handwriting",
        design: Optional[Any] = None
    ) -> bytes:
        if isinstance(pages_text, str):
            pages_list = [pages_text]
        else:
            pages_list = pages_text

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            name="HWTitle",
            parent=styles['Heading1'],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            spaceAfter=12
        )
        body_style = ParagraphStyle(
            name="HWBody",
            parent=styles['Normal'],
            fontName="Helvetica",
            fontSize=11,
            leading=16,
            spaceAfter=8
        )

        story = [
            Paragraph(title, title_style),
            Spacer(1, 10)
        ]

        for p_idx, page_content in enumerate(pages_list):
            if p_idx > 0:
                story.append(Spacer(1, 15))
            for para in page_content.split("\n\n"):
                if para.strip():
                    clean_p = para.replace("\n", " ").strip()
                    story.append(Paragraph(clean_p, body_style))

        doc.build(story)
        return buf.getvalue()

    @staticmethod
    def generate_docx(
        pages_text: Union[List[str], str],
        title: str = "Digitized Handwriting"
    ) -> bytes:
        if isinstance(pages_text, str):
            pages_list = [pages_text]
        else:
            pages_list = pages_text

        doc = Document()
        doc.add_heading(title, level=1)
        for page_content in pages_list:
            for p in page_content.split("\n\n"):
                if p.strip():
                    doc.add_paragraph(p.strip())
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_txt(
        pages_text: Union[List[str], str],
        title: str = "Digitized Handwriting"
    ) -> bytes:
        if isinstance(pages_text, str):
            joined = pages_text
        else:
            joined = "\n\n".join(pages_text)
        content = f"=== {title.upper()} ===\n\n{joined}"
        return content.encode("utf-8")


HandwritingToPdfProcessor = HandwritingProcessor
