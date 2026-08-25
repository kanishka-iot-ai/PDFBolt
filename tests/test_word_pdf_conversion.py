import os
from backend.app.processors.word_to_pdf import WordToPdfProcessor
from pathlib import Path


def test_word_to_pdf_processor_runs(tmp_path):
    # This test only checks the processor flow; requires libreoffice on the runner.
    sample_doc = Path("tests/fixtures/sample.docx")
    if not sample_doc.exists():
        # create a minimal docx fixture
        from docx import Document
        d = Document()
        d.add_heading('Test Doc', level=1)
        d.add_paragraph('This is a sample document for CI validation.')
        sample_dir = sample_doc.parent
        sample_dir.mkdir(parents=True, exist_ok=True)
        d.save(sample_doc)

    proc = WordToPdfProcessor(job_id='test', work_dir=tmp_path, settings={})
    out = proc.process([str(sample_doc)], {})
    assert out.exists()
    assert out.suffix.lower() == '.pdf'
    assert out.stat().st_size > 0
